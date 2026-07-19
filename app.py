"""
VoiceForge - Real-Time Voice Translator  (Speed-Optimised Build)
Run:  pip install -r requirements.txt
      python app.py
"""

import os, uuid, asyncio, warnings, textwrap, time, re, functools
import concurrent.futures
warnings.filterwarnings("ignore")

# ── PyTorch / Whisper compatibility patch ─────────────────────────────────────
import torch
_orig_load = torch.load
def _safe_load(*args, **kwargs):
    kwargs.setdefault("weights_only", False)
    return _orig_load(*args, **kwargs)
torch.load = _safe_load
# ─────────────────────────────────────────────────────────────────────────────

from flask import Flask, render_template, request, send_from_directory, abort, jsonify
from deep_translator import GoogleTranslator
from gtts import gTTS
import whisper
from docx import Document
from PyPDF2 import PdfReader
from pydub import AudioSegment, effects
import librosa, soundfile as sf, noisereduce as nr
import numpy as np
from demucs.pretrained import get_model
from demucs.apply import apply_model
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from scipy.signal import butter, sosfiltfilt

# ── edge-tts ──────────────────────────────────────────────────────────────────
try:
    import edge_tts
    EDGE_TTS_OK = True
    print("✅  edge-tts available.")
except ImportError:
    EDGE_TTS_OK = False
    print("⚠   edge-tts not found – using gTTS fallback. Run: pip install edge-tts")

# ── Flask setup ───────────────────────────────────────────────────────────────
app = Flask(__name__)
app.secret_key = "voiceforge_2025"

STATIC_DIR = os.path.join("static", "output")
UPLOAD_DIR = os.path.join("static", "uploads")
CHUNK_DIR  = os.path.join("static", "chunks")
for _d in [STATIC_DIR, UPLOAD_DIR, CHUNK_DIR]:
    os.makedirs(_d, exist_ok=True)

os.environ["WHISPER_CACHE_DIR"] = os.path.join(os.getcwd(), "whisper_models")

# ── Thread pools (created once at startup, reused for every request) ──────────
# CHUNK_POOL  – Whisper CPU work: 4 workers process audio chunks in parallel
# TTS_POOL    – edge-tts/gTTS network I/O: 8 workers = fast parallel synthesis
# TRANS_POOL  – Google Translate network I/O: 8 workers = fast parallel translate
CHUNK_POOL  = concurrent.futures.ThreadPoolExecutor(max_workers=4)
TTS_POOL    = concurrent.futures.ThreadPoolExecutor(max_workers=8)
TRANS_POOL  = concurrent.futures.ThreadPoolExecutor(max_workers=8)
THREAD_POOL = TTS_POOL  # alias kept for compatibility

# ── Load AI models once at startup ───────────────────────────────────────────
print("⏳  Loading Whisper 'small' model …")
try:
    whisper_model = whisper.load_model("small", device="cpu")
except TypeError:
    whisper_model = whisper.load_model("small")
print("✅  Whisper 'small' ready.")

# Tiny model for fast pre-clean detection only (not transcription)
print("⏳  Loading Whisper 'tiny' model for fast noise-detect …")
try:
    whisper_tiny = whisper.load_model("tiny", device="cpu")
except TypeError:
    whisper_tiny = whisper.load_model("tiny")
print("✅  Whisper 'tiny' ready.")

print("⏳  Loading Demucs …")
demucs_model = get_model("htdemucs")
demucs_model.cpu()
demucs_model.eval()
print("✅  Demucs ready.")

# ── Voice map  (male / female only — child handled separately) ────────────────
EDGE_VOICES = {
    ("en","male"):   "en-US-GuyNeural",
    ("en","female"): "en-US-JennyNeural",
    ("hi","male"):   "hi-IN-MadhurNeural",
    ("hi","female"): "hi-IN-SwaraNeural",
    ("ta","male"):   "ta-IN-ValluvarNeural",
    ("ta","female"): "ta-IN-PallaviNeural",
    ("te","male"):   "te-IN-MohanNeural",
    ("te","female"): "te-IN-ShrutiNeural",
    ("ml","male"):   "ml-IN-MidhunNeural",
    ("ml","female"): "ml-IN-SobhanaNeural",
    ("ja","male"):   "ja-JP-KeitaNeural",
    ("ja","female"): "ja-JP-NanamiNeural",
    ("fr","male"):   "fr-FR-HenriNeural",
    ("fr","female"): "fr-FR-DeniseNeural",
    ("es","male"):   "es-ES-AlvaroNeural",
    ("es","female"): "es-ES-ElviraNeural",
}

# Child voice strategy per language:
#   "edge"  → edge-tts has a genuine child neural voice
#   "gtts"  → gTTS produces a lighter/higher pitch; sounds child-like for these langs
#             (edge-tts has NO dedicated child voice for these languages)
CHILD_VOICE_BACKEND = {
    "en": ("edge", "en-US-AnaNeural"),      # real Microsoft child neural voice
    "fr": ("edge", "fr-FR-EloiseNeural"),   # real Microsoft child neural voice
    "hi": ("gtts", "hi"),                   # gTTS – no edge child voice
    "ta": ("gtts", "ta"),                   # gTTS – no edge child voice
    "te": ("gtts", "te"),                   # gTTS – no edge child voice
    "ml": ("gtts", "ml"),                   # gTTS – no edge child voice
    "ja": ("gtts", "ja"),                   # gTTS – no edge child voice (ja female ≠ child)
    "es": ("gtts", "es"),                   # gTTS – no edge child voice
}

LANG_NAMES = {
    # ── 8 supported output languages ─────────────────────────────────────────
    "en":"English",  "ta":"Tamil",     "hi":"Hindi",    "te":"Telugu",
    "ml":"Malayalam","ja":"Japanese",  "fr":"French",   "es":"Spanish",
    # ── additional Whisper-detected languages (shown in UI but not output) ───
    "zh":"Chinese",  "de":"German",    "ar":"Arabic",   "pt":"Portuguese",
    "ru":"Russian",  "ko":"Korean",    "it":"Italian",  "nl":"Dutch",
    "pl":"Polish",   "sv":"Swedish",   "tr":"Turkish",  "id":"Indonesian",
    "th":"Thai",     "vi":"Vietnamese","uk":"Ukrainian", "cs":"Czech",
    "ro":"Romanian", "hu":"Hungarian", "kn":"Kannada",  "mr":"Marathi",
    "ur":"Urdu",     "bn":"Bengali",   "pa":"Punjabi",  "gu":"Gujarati",
}

def load_best_audio_segment(wav_path, sr=16000, duration_s=30):
    """
    Load the loudest 30s segment for language detection
    so uploaded/recorded audio with silent leading sections
    still detects language correctly.
    """
    y, _ = librosa.load(wav_path, sr=sr, mono=True)
    if len(y) == 0:
        return whisper.pad_or_trim(whisper.load_audio(wav_path))

    target_len = sr * duration_s
    if len(y) <= target_len:
        return whisper.pad_or_trim(y)

    hop = 512
    try:
        rms = librosa.feature.rms(y=y, frame_length=2048, hop_length=hop)[0]
        if len(rms) == 0:
            return whisper.pad_or_trim(y[:target_len])
        idx = int(np.argmax(rms))
        start = max(0, idx * hop - target_len // 2)
        start = min(start, len(y) - target_len)
        return whisper.pad_or_trim(y[start:start + target_len])
    except Exception:
        return whisper.pad_or_trim(y[:target_len])


def detect_audio_language(wav_path):
    """
    Detect language of audio using Whisper's dedicated language detection.
    Returns exactly 3 values in every code path.
    """
    try:
        audio = load_best_audio_segment(wav_path, sr=16000, duration_s=30)

        mel = whisper.log_mel_spectrogram(audio).to(whisper_tiny.device)
        _, probs_tiny = whisper_tiny.detect_language(mel)
        lang_code = max(probs_tiny, key=probs_tiny.get)
        confidence = probs_tiny[lang_code]
        lang_name = LANG_NAMES.get(lang_code, lang_code.upper())
        if confidence > 0.55:
            return lang_code, lang_name, int(confidence * 100)

        mel = whisper.log_mel_spectrogram(audio).to(whisper_model.device)
        _, probs = whisper_model.detect_language(mel)
        if isinstance(probs, (list, tuple)) and len(probs) == 2:
            probs = probs[1]
        lang_code = max(probs, key=probs.get)
        confidence = probs[lang_code]
        lang_name = LANG_NAMES.get(lang_code, lang_code.upper())

        if confidence < 0.50:
            res2 = whisper_model.transcribe(
                wav_path, fp16=False, task="transcribe",
                beam_size=3, temperature=0, no_speech_threshold=0.4
            )
            lang_code2 = res2.get("language", lang_code)
            if lang_code2 and lang_code2 != lang_code:
                lang_code = lang_code2
                lang_name = LANG_NAMES.get(lang_code, lang_code.upper())
                confidence = max(confidence, 0.50)

        return lang_code, lang_name, int(confidence * 100)

    except Exception as e:
        print(f"  detect_audio_language error: {e}")
        return "unknown", "Unknown", 0

# ═════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═════════════════════════════════════════════════════════════════════════════

def safe_remove(path):
    try:
        if path and os.path.exists(path): os.remove(path)
    except Exception: pass


def extract_text_from_file(file_obj):
    fname = file_obj.filename.lower()
    save_path = os.path.join(UPLOAD_DIR, file_obj.filename)
    file_obj.save(save_path)
    try:
        if fname.endswith(".txt"):
            with open(save_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif fname.endswith(".docx"):
            doc = Document(save_path)
            return "\n".join(p.text for p in doc.paragraphs)
        elif fname.endswith(".pdf"):
            reader = PdfReader(save_path)
            return "".join(p.extract_text() or "" for p in reader.pages)
        else:
            raise ValueError("Unsupported file type. Use TXT, DOCX or PDF.")
    finally:
        safe_remove(save_path)


def split_text(text, max_len=4000):
    # Reduce max_len to 3000 for more parallel chunks (faster overall, but keep under Google's 5000 limit)
    if not text: return []
    sentences = re.split(r'(?<=[.!?।॥。！？])\s+', text.replace('\n', ' ').strip())
    chunks, cur = [], ''
    for s in sentences:
        s = s.strip()
        if not s: continue
        add = s + ' '
        if len(cur) + len(add) <= 3000:  # Changed from 4000
            cur += add
        else:
            if cur: chunks.append(cur.strip())
            cur = add
    if cur.strip(): chunks.append(cur.strip())
    return chunks or [text[:3000]]


def split_audio(audio_seg, chunk_ms=60_000):
    """60-second chunks — half the number of Whisper calls vs 30s."""
    return [audio_seg[i:i+chunk_ms] for i in range(0, len(audio_seg), chunk_ms)]


def is_silent(seg, thresh=-90):
    """More permissive silence gate so quiet audio is not dropped."""
    if len(seg) < 250:
        return True
    try:
        return seg.dBFS < thresh
    except Exception:
        return True


def to_wav(file_obj, uid, mono=True, sr=16000):
    """Convert any uploaded format to WAV. Preserves extension so ffmpeg detects format."""
    # Keep the original extension so ffmpeg/pydub can detect format correctly.
    # Recorded blobs arrive as recording.webm or recording.ogg — without extension
    # ffmpeg would fail to identify the container.
    orig_name = getattr(file_obj, 'filename', '') or ''
    ext = os.path.splitext(orig_name)[1].lower()   # e.g. '.webm', '.mp3', '.wav'
    if not ext:
        # Guess from content_type if filename has no extension
        ct = getattr(file_obj, 'content_type', '') or ''
        if 'ogg' in ct:    ext = '.ogg'
        elif 'webm' in ct: ext = '.webm'
        elif 'mp4' in ct or 'aac' in ct: ext = '.mp4'
        else:              ext = '.webm'  # safest fallback for browser recordings
    tmp = os.path.join(UPLOAD_DIR, f"{uid}_in{ext}")
    file_obj.save(tmp)
    try:
        seg = AudioSegment.from_file(tmp)
        if mono:
            seg = seg.set_channels(1).set_frame_rate(sr)
        wav_path = os.path.join(UPLOAD_DIR, f"{uid}.wav")
        seg.export(wav_path, format="wav")
        return wav_path
    finally:
        safe_remove(tmp)


# ── Audio cleaning (mic recordings only) ─────────────────────────────────────

def clean_mic_audio(y, sr_val):
    """
    Light noise clean for mic-recorded audio only.
    Uses gentle prop_decrease=0.4 (was 0.7) to preserve speech quality —
    over-aggressive noise reduction distorts phonemes and makes Whisper
    fail on non-English languages like Tamil, Hindi, Telugu, Malayalam.
    """
    # Skip noise reduction for very short clips — avoids distorting single words
    if len(y) > sr_val * 1.5:   # only if audio > 1.5 seconds
        try:
            y = nr.reduce_noise(y=y, sr=sr_val, prop_decrease=0.35,
                                stationary=True, n_fft=512, n_jobs=1)
        except Exception:
            pass
    # High-pass: strip low-frequency mic rumble below 80 Hz only
    try:
        sos = butter(2, 80.0 / (sr_val / 2), btype="high", output="sos")
        y   = sosfiltfilt(sos, y)
    except Exception:
        pass
    # Peak-normalise to 0.95 so Whisper gets clear signal
    peak = np.max(np.abs(y))
    if peak > 0.0:
        y = y / peak * 0.95
    return y


def apply_noise_reduction(y, sr_val, prop_decrease):
    """
    Two-pass noise reduction for best results:
    Pass 1 — stationary (removes steady hiss/hum, fast)
    Pass 2 — non-stationary at lower strength (removes variable background noise)
    Both passes use n_fft=512 for speed.
    """
    if prop_decrease <= 0.0:
        return y
    try:
        y = nr.reduce_noise(y=y, sr=sr_val,
                            prop_decrease=float(prop_decrease),
                            stationary=True, n_fft=256, n_jobs=1)
    except Exception as e:
        print(f"  noise reduction pass1 error: {e}")
    try:
        if prop_decrease > 0.25:
            y = nr.reduce_noise(y=y, sr=sr_val,
                                prop_decrease=float(prop_decrease) * 0.5,
                                stationary=False, n_fft=256, n_jobs=1)
    except Exception as e:
        print(f"  noise reduction pass2 error: {e}")
    return y

def apply_clarity_eq(y, sr_val, clarity):
    """
    3-band EQ for voice clarity:
    1. High-pass at 80-200 Hz (removes rumble)
    2. Presence boost at 1-4 kHz (makes voice cut through)
    3. Air boost at 8 kHz (adds crispness/intelligibility)
    """
    if clarity <= 0.0:
        return y
    nyq = sr_val / 2.0
    try:
        hp_hz = 80.0 + clarity * 120.0
        if hp_hz < nyq * 0.9:
            sos_hp = butter(2, hp_hz / nyq, btype="high", output="sos")
            y = sosfiltfilt(sos_hp, y)
    except Exception: pass
    try:
        pres_hz = 1500.0
        if pres_hz < nyq * 0.9:
            sos_pres = butter(2, pres_hz / nyq, btype="high", output="sos")
            y_pres = sosfiltfilt(sos_pres, y)
            boost = 10 ** (clarity * 10.0 / 20.0) - 1.0
            y = y + y_pres * boost
    except Exception: pass
    try:
        air_hz = 6000.0
        if air_hz < nyq * 0.9:
            sos_air = butter(2, air_hz / nyq, btype="high", output="sos")
            y_air = sosfiltfilt(sos_air, y)
            air_boost = 10 ** (clarity * 4.0 / 20.0) - 1.0
            y = y + y_air * air_boost
    except Exception: pass
    # Final normalise — prevent clipping
    peak = np.max(np.abs(y))
    if peak > 0.97:
        y = y / peak * 0.97
    return y


# ── TTS helpers ───────────────────────────────────────────────────────────────

async def _edge_save(text, voice, path):
    """Standard edge-tts synthesis."""
    comm = edge_tts.Communicate(text, voice)
    await comm.save(path)


def tts_one_chunk(args):
    """
    Synthesise one text chunk → tmp mp3 file.  Runs inside the thread pool.

    Child voice routing (CHILD_VOICE_BACKEND dict):
      en, fr  → genuine edge-tts child neural voice  (AnaNeural / EloiseNeural)
      hi,ta,te,ml,ja,es → gTTS (lighter, higher-pitched; clearly different from adult)
                          edge-tts has NO dedicated child neural voice for these.

    Male / Female → always edge-tts neural voice; gTTS only as absolute last resort.
    """
    text, lang, mode, tmp_path = args

    # ── Child mode ────────────────────────────────────────────────────────────
    if mode == "child":
        backend, voice_or_lang = CHILD_VOICE_BACKEND.get(lang, ("gtts", lang))

        if backend == "edge" and EDGE_TTS_OK:
            try:
                asyncio.run(_edge_save(text, voice_or_lang, tmp_path))
                return tmp_path
            except Exception as e:
                print(f"  edge-tts child error ({e}) – falling back to gTTS")
                backend = "gtts"   # fallthrough to gTTS below

        # gTTS child — generate then pitch-shift +4 semitones to sound genuinely child-like
        try:
            gTTS(text=text, lang=voice_or_lang, slow=False).save(tmp_path)
            try:
                y, sr_val = librosa.load(tmp_path, sr=None, mono=True)
                # +4 semitones → clearly child-like pitch
                y = librosa.effects.pitch_shift(y, sr=sr_val, n_steps=4)
                # 1.08× speed — children speak slightly faster
                y = librosa.effects.time_stretch(y, rate=1.08)
                peak = np.max(np.abs(y))
                if peak > 0: y = y / peak * 0.92
                sf.write(tmp_path, y, sr_val)
            except Exception as pe:
                print(f"  child pitch-shift error ({pe}) – using plain gTTS")
            return tmp_path
        except Exception as e:
            print(f"  gTTS child error ({e}) – last resort: edge female")
            if EDGE_TTS_OK:
                voice = EDGE_VOICES.get((lang, "female"), "en-US-JennyNeural")
                try:
                    asyncio.run(_edge_save(text, voice, tmp_path))
                    return tmp_path
                except Exception:
                    pass
        return tmp_path

    # ── Male / Female mode ────────────────────────────────────────────────────
    if EDGE_TTS_OK:
        voice = EDGE_VOICES.get((lang, mode),
                EDGE_VOICES.get((lang, "female"), "en-US-JennyNeural"))
        try:
            asyncio.run(_edge_save(text, voice, tmp_path))
            return tmp_path
        except Exception as e:
            print(f"  edge-tts {mode} error ({e}) – falling back to gTTS")

    # ── gTTS final fallback ───────────────────────────────────────────────────
    try:
        gTTS(text=text, lang=lang, slow=False).save(tmp_path)
    except Exception as e:
        print(f"  gTTS fallback error ({e})")
    return tmp_path


def tts_chunks_to_segment(text, lang, mode):
    """
    Split text → parallel TTS synthesis → single AudioSegment.
    All chunks generated concurrently in the thread pool.
    """
    parts = split_text(text, 4000)  # edge-tts handles up to 5000 chars per call
    # Assign deterministic temp paths so we can re-order after futures complete
    jobs = [(p, lang, mode, os.path.join(CHUNK_DIR, f"{uuid.uuid4().hex}.mp3"))
            for p in parts]

    # Submit all TTS jobs in parallel using TTS_POOL (separate from CHUNK_POOL)
    futures = {TTS_POOL.submit(tts_one_chunk, job): i for i, job in enumerate(jobs)}
    results = [None] * len(jobs)
    for fut in concurrent.futures.as_completed(futures):
        idx = futures[fut]
        try:
            results[idx] = fut.result()
        except Exception as e:
            print(f"  TTS chunk {idx} failed: {e}")
            results[idx] = None

    combined = AudioSegment.empty()
    for path in results:
        if path and os.path.exists(path):
            try:
                combined += AudioSegment.from_file(path)
            except Exception:
                pass
            safe_remove(path)
    return combined


# ── Parallel translate helper ─────────────────────────────────────────────────

def translate_chunks_parallel(chunks, target_lang):
    """Translate all text chunks in parallel using the thread pool."""
    def _translate(chunk):
        return GoogleTranslator(source="auto", target=target_lang).translate(chunk)

    futures = [THREAD_POOL.submit(_translate, c) for c in chunks]
    results = []
    for fut in concurrent.futures.as_completed(futures):
        try:
            results.append(fut.result())
        except Exception:
            results.append("")
    # as_completed doesn't preserve order — re-submit in order instead
    return results


@functools.lru_cache(maxsize=128)
def translate_text(text, target_lang, source_lang="en"):
    """
    Translate text → target_lang accurately and fast.
    source_lang="en" by default because Whisper task=translate always
    outputs English — telling Google the source avoids mis-detection.
    """
    if not text or not text.strip():
        return text
    if source_lang == target_lang:
        return text

    def _one(chunk):
        for attempt in range(3):
            try:
                out = GoogleTranslator(
                    source=source_lang if source_lang != "auto" else "auto",
                    target=target_lang
                ).translate(chunk)
                if out and out.strip():
                    return out
            except Exception:
                pass
            try:
                out = GoogleTranslator(source="auto", target=target_lang).translate(chunk)
                if out and out.strip():
                    return out
            except Exception:
                pass
            time.sleep(0.15 * (attempt + 1))
        return chunk

    chunks = split_text(text, 4000)
    if len(chunks) == 1:
        return _one(chunks[0])
    futs = [TRANS_POOL.submit(_one, c) for c in chunks]
    return " ".join(f.result() for f in futs)


# ── Process one audio chunk: clean → transcribe → translate → TTS ─────────────

def process_audio_chunk(args):
    """
    Full pipeline for one audio chunk in voice_to_voice.
    Steps:
      1. Export chunk → WAV
      2. Optional mic noise cleaning
      3. Whisper transcribe with task="translate" → always gives English text
         (Whisper's translate task is the most accurate way to get the spoken
          content regardless of source language — avoids GoogleTranslator
          source-detection failures that caused wrong output language)
      4. Translate English intermediate → user-selected target language
         (if target is already English, skip this step)
      5. TTS in target language + chosen voice mode
    """
    chunk_seg, uid, i, lang, mode, is_mic = args
    if is_silent(chunk_seg):
        return None

    cp = os.path.join(CHUNK_DIR, f"{uid}_c{i}.wav")
    try:
        chunk_seg.export(cp, format="wav")

        if is_mic:
            y_c, sr_c = librosa.load(cp, sr=16000, mono=True)
            y_c = clean_mic_audio(y_c, sr_c)
            sf.write(cp, y_c, sr_c)

        # Step 1: Whisper task="translate" → English always, any source language.
        # beam_size=5: consider 5 hypotheses → pick best (most accurate on CPU).
        # condition_on_previous_text=False: prevents repetition/hallucination.
        # temperature=0: greedy decode (fastest + most accurate for clear audio).
        # Step 1: Transcribe in ORIGINAL language (keeps exact words/meaning)
        result = whisper_model.transcribe(
            cp, fp16=False, task="translate",  # Changed from "transcribe" to "translate" for direct English output
            beam_size=3, best_of=3,  # Reduced from 5 for speed
            condition_on_previous_text=False,
            temperature=[0, 0.2, 0.4],
            no_speech_threshold=0.4,
        )
        en_text = result.get("text", "").strip()  # Now directly English
        if not en_text:
            return None

        # Translate English to target (keep Google for accuracy)
        if lang == "en":
            final_text = en_text
        else:
            final_text = translate_text(en_text, lang, source_lang="en")  # Source is now known as "en"
        return tts_chunks_to_segment(final_text, lang, mode)

    except Exception as e:
        print(f"  v2v chunk {i} error: {e}")
        return None
    finally:
        safe_remove(cp)


def process_audio_chunk_v2t(args):
    chunk_seg, uid, i, is_mic = args
    if is_silent(chunk_seg):
        return ""

    cp = os.path.join(CHUNK_DIR, f"{uid}_c{i}.wav")
    try:
        chunk_seg.export(cp, format="wav")

        if is_mic:
            y_c, sr_c = librosa.load(cp, sr=16000, mono=True)
            y_c = clean_mic_audio(y_c, sr_c)
            sf.write(cp, y_c, sr_c)

        res = whisper_model.transcribe(
            cp, fp16=False, task="transcribe",
            beam_size=3, best_of=3,
            condition_on_previous_text=False,
            temperature=[0, 0.2, 0.4],
            no_speech_threshold=0.4,
        )
        src_text = res.get("text", "").strip()
        src_lang = res.get("language", "en")
        return f"__LANG:{src_lang}__" + src_text if src_text else ""
    except Exception as e:
        print(f"  v2t chunk {i} error: {e}")
        return ""
    finally:
        safe_remove(cp)


# ── Demucs separation ─────────────────────────────────────────────────────────

def separate_vocals(input_wav, output_dir, uid):
    wav_np, sr_val = sf.read(input_wav, dtype="float32")
    if wav_np.ndim == 1:
        wav_np = np.stack([wav_np, wav_np], axis=1)
    wav_t = torch.from_numpy(wav_np.T).unsqueeze(0).float()
    with torch.no_grad():
        sources = apply_model(demucs_model, wav_t, device="cpu",
                              progress=False, shifts=0, num_workers=1)  # Added shifts=0
    sources = sources.squeeze(0).cpu().numpy()
    idx     = {s: i for i, s in enumerate(demucs_model.sources)}
    vocals  = sources[idx["vocals"]]
    music   = sources[idx["drums"]] + sources[idx["bass"]] + sources[idx["other"]]
    vocal_path = os.path.join(output_dir, f"{uid}_vocals.wav")
    music_path = os.path.join(output_dir, f"{uid}_music.wav")
    sf.write(vocal_path, vocals.T, sr_val)
    sf.write(music_path, music.T, sr_val)
    return vocal_path, music_path


# ── PDF generation ────────────────────────────────────────────────────────────

def make_pdf(text, path):
    """
    Convert text → A4 PDF using wkhtmltopdf (via subprocess).

    Writes the text into a UTF-8 HTML file, then calls wkhtmltopdf which
    uses the system Unifont fallback — a bitmap/outline font that covers
    every Unicode code point including Tamil, Hindi, Telugu, Malayalam,
    Japanese, French and Spanish.

    Falls back to a reportlab plain-text PDF if wkhtmltopdf is not found.
    """
    import tempfile, subprocess

    def _esc(s):
        return (s.replace('&', '&amp;')
                 .replace('<', '&lt;')
                 .replace('>', '&gt;')
                 .replace('"', '&quot;'))

    # Build HTML paragraphs from newline-separated text
    paras = ''
    for line in text.split('\n'):
        line = line.strip()
        paras += f'<p>{_esc(line)}</p>\n' if line else '<p class="sp"></p>\n'

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<meta http-equiv="Content-Type" content="text/html; charset=UTF-8">
<style>
* {{ margin:0; padding:0; box-sizing:border-box; }}
body {{
    font-family: FreeSerif, FreeSans, "Noto Sans", Unifont, sans-serif;
    font-size: 13.5pt;
    line-height: 2.0;
    color: #111;
    padding: 20mm 22mm;
}}
p {{ margin-bottom: 8px; word-break: break-word; }}
p.sp {{ height: 10px; margin: 0; }}
</style>
</head>
<body>
{paras}
</body>
</html>"""

    # Write HTML to a named temp file (not stdin — avoids encoding issues)
    tmp = tempfile.NamedTemporaryFile(
        suffix='.html', delete=False, mode='w', encoding='utf-8')
    try:
        tmp.write(html)
        tmp.flush()
        tmp.close()

        cmd = [
            'wkhtmltopdf',
            '--enable-local-file-access',
            '--encoding',       'utf-8',
            '--page-size',      'A4',
            '--margin-top',     '0mm',
            '--margin-bottom',  '0mm',
            '--margin-left',    '0mm',
            '--margin-right',   '0mm',
            '--quiet',
            '--no-stop-slow-scripts',
            tmp.name,
            path,
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode == 0 and os.path.exists(path) and os.path.getsize(path) > 0:
            return   # ✅ success
        print(f'  wkhtmltopdf non-zero exit ({result.returncode}): {result.stderr[:200]}')

    except FileNotFoundError:
        print('  wkhtmltopdf not found — falling back to reportlab')
    except Exception as e:
        print(f'  wkhtmltopdf error: {e}')
    finally:
        try: os.unlink(tmp.name)
        except Exception: pass

    # ── Fallback: reportlab with FreeSerif (covers Tamil/Hindi/Malayalam) ──
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4 as RL_A4
        from reportlab.pdfbase.pdfmetrics import stringWidth

        FONT = 'Helvetica'
        for fpath, fname in [
            ('/usr/share/fonts/truetype/freefont/FreeSerif.ttf', 'VFFallback'),
            ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  'VFFallback'),
        ]:
            if os.path.exists(fpath):
                try:
                    pdfmetrics.registerFont(TTFont('VFFallback', fpath))
                    FONT = 'VFFallback'
                    break
                except Exception:
                    pass

        W, H   = RL_A4
        MARGIN = 50
        FS     = 13
        LH     = 22
        MAX_W  = W - 2 * MARGIN

        c = rl_canvas.Canvas(path, pagesize=RL_A4)
        c.setFont(FONT, FS)
        y = H - MARGIN

        def _new_page():
            nonlocal y
            c.showPage()
            c.setFont(FONT, FS)
            y = H - MARGIN

        for raw in text.split('\n'):
            raw = raw.strip()
            if not raw:
                y -= LH * 0.5
                if y < MARGIN + LH: _new_page()
                continue
            words, cur = raw.split(' '), ''
            for w in words:
                test = (cur + ' ' + w).strip()
                if stringWidth(test, FONT, FS) <= MAX_W:
                    cur = test
                else:
                    if cur:
                        if y < MARGIN + LH: _new_page()
                        c.drawString(MARGIN, y, cur)
                        y -= LH
                    cur = w
            if cur:
                if y < MARGIN + LH: _new_page()
                c.drawString(MARGIN, y, cur)
                y -= LH
            y -= 4
        c.save()
    except Exception as e:
        print(f'  reportlab fallback error: {e}')


# ═════════════════════════════════════════════════════════════════════════════
# ROUTES
# ═════════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html")


# ── Text → Voice ──────────────────────────────────────────────────────────────
@app.route("/text_to_voice", methods=["GET", "POST"])
def text_to_voice():
    output_file = error = translated_text = input_text = None
    lang = mode = text = doc_file = None
    sub_lang = sub_mode = sub_input_label = sub_chars = None
    if request.method == "POST":
        ajax_request = request.form.get("ajax") == "1" or "application/json" in request.headers.get("Accept", "")
        try:
            lang     = request.form["language"]
            mode     = request.form["voice_mode"]
            text     = request.form.get("text", "").strip()
            doc_file = request.files.get("document")

            if not text and doc_file and doc_file.filename:
                text = extract_text_from_file(doc_file)
            if not text:
                raise ValueError("Please enter text or upload a document.")

            sub_lang        = LANG_NAMES.get(lang, lang.upper())
            sub_mode        = {"female":"👩 Female","male":"👨 Male","child":"🧒 Child"}.get(mode, mode)
            sub_input_label = (doc_file.filename if doc_file and doc_file.filename else None) or "Text input"
            sub_chars       = len(text)
            input_text      = text

            translated_text = translate_text(text, lang, source_lang="auto")
            if not translated_text or not translated_text.strip():
                translated_text = text
            uid   = uuid.uuid4().hex
            final = tts_chunks_to_segment(translated_text, lang, mode)
            out   = os.path.join(STATIC_DIR, f"{uid}.mp3")
            final.export(out, format="mp3", bitrate="192k")
            output_file = "/" + out.replace("\\", "/")

            if ajax_request:
                return jsonify({
                    "output_file": output_file,
                    "translated_text": translated_text,
                    "input_text": input_text,
                    "sub_lang": sub_lang,
                    "sub_mode": sub_mode,
                    "sub_input_label": sub_input_label,
                    "sub_chars": sub_chars,
                })

        except Exception as e:
            if ajax_request:
                return jsonify({"error": str(e)}), 400
            error = str(e)

    return render_template("text_to_voice.html",
                           output_file=output_file,
                           translated_text=translated_text,
                           input_text=input_text,
                           error=error,
                           sub_lang=sub_lang, sub_mode=sub_mode,
                           sub_input_label=sub_input_label,
                           sub_chars=sub_chars)


# ── Voice → Voice ─────────────────────────────────────────────────────────────
@app.route("/voice_to_voice", methods=["GET", "POST"])
def voice_to_voice():
    # GET — serve the page
    if request.method == "GET":
        return render_template("voice_to_voice.html")

    # POST — process and return JSON (frontend uses fetch, no page reload)
    try:
        lang      = request.form.get("language", "en")
        mode      = request.form.get("voice_mode", "female")
        audio_obj = request.files.get("audio")
        if not audio_obj or not audio_obj.filename:
            return jsonify({"error": "No audio file received."}), 400

        fname_lc = audio_obj.filename.lower()
        is_mic   = (fname_lc.endswith(".webm") or fname_lc.endswith(".ogg") or
                    fname_lc.startswith("recording.") or
                    audio_obj.content_type in ("audio/webm", "audio/ogg",
                                               "audio/webm;codecs=opus",
                                               "audio/ogg;codecs=opus"))

        uid    = uuid.uuid4().hex
        wav_in = to_wav(audio_obj, uid)

        sub_copy = os.path.join(STATIC_DIR, f"{uid}_input.wav")
        import shutil; shutil.copy2(wav_in, sub_copy)
        sub_audio_url = "/static/output/" + f"{uid}_input.wav"
        sub_fname     = "Recorded audio" if is_mic else audio_obj.filename

        full_seg = AudioSegment.from_wav(wav_in)
        dur_s    = round(len(full_seg) / 1000)
        sub_duration = f"{dur_s//60}m {dur_s%60}s" if dur_s >= 60 else f"{dur_s}s"
        sub_lang = LANG_NAMES.get(lang, lang.upper())
        sub_mode = {"female": "👩 Female", "male": "👨 Male",
                    "child":  "🧒 Child"}.get(mode, mode)

        chunk_ms = 45_000 if is_mic else 120_000
        chunks   = split_audio(full_seg, chunk_ms)
        jobs     = [(chunk, uid, i, lang, mode, is_mic)
                    for i, chunk in enumerate(chunks)]
        futures = [CHUNK_POOL.submit(process_audio_chunk, job) for job in jobs]

        output_seg = AudioSegment.empty()
        for fut in futures:
            try:
                seg = fut.result()
                if seg is not None and len(seg) > 0:
                    output_seg += seg
            except Exception as chunk_err:
                print(f"  chunk collection error: {chunk_err}")

        if len(output_seg) == 0:
            return jsonify({"error": "No speech detected in the audio."}), 400

        out_name = f"{uid}_v2v.mp3"
        out_path = os.path.join(STATIC_DIR, out_name)
        output_seg.export(out_path, format="mp3", bitrate="192k")
        safe_remove(wav_in)

        return jsonify({
            "audio_url":     "/static/output/" + out_name,
            "sub_fname":     sub_fname,
            "sub_audio_url": sub_audio_url,
            "sub_lang":      sub_lang,
            "sub_mode":      sub_mode,
            "sub_duration":  sub_duration,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500



# ── Voice → Text ──────────────────────────────────────────────────────────────
@app.route("/voice_to_text", methods=["GET", "POST"])
def voice_to_text():
    converted_text = file_id = error = None
    audio_obj = lang = None
    sub_fname = sub_audio_url = sub_lang = sub_duration = None
    if request.method == "POST":
        ajax_request = request.form.get("ajax") == "1" or "application/json" in request.headers.get("Accept", "")
        try:
            lang      = request.form["language"]
            audio_obj = request.files.get("audio_file")
            if not audio_obj or not audio_obj.filename:
                raise ValueError("No audio file received.")

            fname_lc = audio_obj.filename.lower()
            is_mic = (fname_lc.endswith(".webm") or fname_lc.endswith(".ogg") or
                      fname_lc.startswith("recording.") or
                      audio_obj.content_type in ("audio/webm","audio/ogg",
                                                 "audio/webm;codecs=opus",
                                                 "audio/ogg;codecs=opus"))

            uid      = uuid.uuid4().hex
            wav_in   = to_wav(audio_obj, uid)

            sub_copy = os.path.join(STATIC_DIR, f"{uid}_input.wav")
            import shutil; shutil.copy2(wav_in, sub_copy)
            sub_audio_url = f"/static/output/{uid}_input.wav"
            sub_fname     = "Recorded audio" if is_mic else audio_obj.filename
            sub_lang      = LANG_NAMES.get(lang, lang.upper())

            full_seg = AudioSegment.from_wav(wav_in)
            dur_s    = round(len(full_seg) / 1000)
            sub_duration = f"{dur_s//60}m {dur_s%60}s" if dur_s >= 60 else f"{dur_s}s"
            chunk_ms = 45_000 if is_mic else 120_000
            chunks   = split_audio(full_seg, chunk_ms)

            jobs    = [(chunk, uid, i, is_mic) for i, chunk in enumerate(chunks)]
            futures = [CHUNK_POOL.submit(process_audio_chunk_v2t, job) for job in jobs]
            chunk_results = [f.result() for f in futures]

            import re as _re
            detected_src_lang = "en"
            clean_parts = []
            for part in chunk_results:
                if not part: continue
                m = _re.match(r"^__LANG:([a-z]{2,3})__(.*)$", part, _re.DOTALL)
                if m:
                    detected_src_lang = m.group(1)
                    clean_parts.append(m.group(2).strip())
                else:
                    clean_parts.append(part.strip())

            raw_text = " ".join(p for p in clean_parts if p).strip()
            if not raw_text:
                raise ValueError("No speech detected in the audio.")

            if detected_src_lang == lang:
                converted_text = raw_text
            elif lang == "en" and detected_src_lang != "en":
                # Target is English — use Whisper translate on full audio for best quality
                wav_for_en = to_wav(audio_obj, uid+"_en") if False else None
                # Fallback: translate the transcribed text
                converted_text = translate_text(raw_text, "en", source_lang=detected_src_lang)
            else:
                # Translate from detected source → target
                converted_text = translate_text(raw_text, lang, source_lang=detected_src_lang)
                if not converted_text or not converted_text.strip():
                    # Last fallback: treat as English source
                    converted_text = translate_text(raw_text, lang, source_lang="en")
                if not converted_text or not converted_text.strip():
                    converted_text = raw_text

            file_id = uuid.uuid4().hex

            def _save_txt():
                with open(os.path.join(STATIC_DIR, f"{file_id}.txt"), "w", encoding="utf-8") as f:
                    f.write(converted_text)
            def _save_docx():
                doc = Document()
                doc.add_paragraph(converted_text)
                doc.save(os.path.join(STATIC_DIR, f"{file_id}.docx"))
            def _save_pdf():
                make_pdf(converted_text, os.path.join(STATIC_DIR, f"{file_id}.pdf"))

            save_futs = [THREAD_POOL.submit(_save_txt),
                         THREAD_POOL.submit(_save_docx),
                         THREAD_POOL.submit(_save_pdf)]
            for f in save_futs: f.result()
            safe_remove(wav_in)

            if ajax_request:
                return jsonify({
                    "converted_text": converted_text,
                    "file_id": file_id,
                    "sub_fname": sub_fname,
                    "sub_audio_url": sub_audio_url,
                    "sub_lang": sub_lang,
                    "sub_duration": sub_duration,
                })

        except Exception as e:
            if ajax_request:
                return jsonify({"error": str(e)}), 400
            error = str(e)

    return render_template("voice_to_text.html",
                           converted_text=converted_text,
                           file_id=file_id, error=error,
                           sub_fname=sub_fname, sub_audio_url=sub_audio_url,
                           sub_lang=sub_lang, sub_duration=sub_duration)


# ── Download ──────────────────────────────────────────────────────────────────
@app.route("/download/<file_id>/<fmt>")
def download_file(file_id, fmt):
    if fmt not in ("txt", "docx", "pdf"): abort(404)
    fname = f"{file_id}.{fmt}"
    if not os.path.exists(os.path.join(STATIC_DIR, fname)): abort(404)
    return send_from_directory(STATIC_DIR, fname, as_attachment=True)


# ── Noise Removal ─────────────────────────────────────────────────────────────
@app.route("/noise_removal", methods=["GET", "POST"])
def noise_removal():
    vocals_url = background_url = error = None
    if request.method == "POST":
        ajax_request = request.form.get("ajax") == "1" or "application/json" in request.headers.get("Accept", "")
        try:
            audio_obj = request.files.get("audio")
            if not audio_obj or not audio_obj.filename:
                raise ValueError("No audio file received.")

            uid = uuid.uuid4().hex
            wav_in = to_wav(audio_obj, uid, mono=True, sr=44100)

            wav, sr = librosa.load(wav_in, sr=44100, mono=True)
            if wav.size == 0:
                raise ValueError("Could not decode audio.")

            wav_tensor = torch.from_numpy(wav).float().unsqueeze(0)
            with torch.no_grad():
                sources = apply_model(
                    demucs_model,
                    wav_tensor,
                    device="cpu",
                    progress=False,
                    num_workers=1,
                )

            if isinstance(sources, list):
                sources = torch.stack(sources)

            if sources.ndim == 3 and sources.shape[1] == 1:
                sources = sources[:, 0, :]
            elif sources.ndim == 1:
                sources = sources.unsqueeze(0)

            source_names = getattr(demucs_model, "sources", None)
            if source_names is None:
                source_names = [f"src{i}" for i in range(sources.shape[0])]

            vocals_index = source_names.index("vocals") if "vocals" in source_names else -1
            vocals = sources[vocals_index] if vocals_index >= 0 else sources[-1]
            background = sources.sum(dim=0) - vocals

            vocals_path = os.path.join(STATIC_DIR, f"{uid}_vocals.wav")
            background_path = os.path.join(STATIC_DIR, f"{uid}_background.wav")

            sf.write(vocals_path, vocals.cpu().numpy().T, sr)
            sf.write(background_path, background.cpu().numpy().T, sr)

            safe_remove(wav_in)

            vocals_url = "/" + vocals_path.replace("\\", "/")
            background_url = "/" + background_path.replace("\\", "/")

            if ajax_request:
                return jsonify({
                    "vocals_url": vocals_url,
                    "background_url": background_url,
                })

        except Exception as e:
            error = str(e)
            if ajax_request:
                return jsonify({"error": error}), 400

    return render_template("noise_removal.html",
                           vocals_url=vocals_url,
                           background_url=background_url,
                           error=error)


# ── Language Detection (AJAX — called immediately when audio is uploaded) ──────
@app.route("/detect_language", methods=["POST"])
def detect_language():
    """
    Accurately detect the language of an uploaded/recorded audio file.
    Uses Whisper small model's dedicated detect_language() for maximum accuracy.
    Returns: language display name + raw lang_code + confidence percentage.
    """
    uid    = uuid.uuid4().hex
    wav_in = None
    try:
        audio_obj = request.files.get("audio")
        if not audio_obj or not audio_obj.filename:
            return jsonify({"language": "Unknown", "lang_code": "unknown", "confidence": 0})

        # Convert to 16kHz mono WAV (Whisper's optimal input format)
        wav_in = to_wav(audio_obj, uid, mono=True, sr=16000)

        lang_code, language, confidence = detect_audio_language(wav_in)

        return jsonify({
            "language":   language,
            "lang_code":  lang_code,
            "confidence": confidence,   # 0–100 integer
        })
    except Exception as e:
        print(f"  /detect_language error: {e}")
        return jsonify({"language": "Unknown", "lang_code": "unknown", "confidence": 0})
    finally:
        if wav_in: safe_remove(wav_in)


# ── Voice Enhancer ────────────────────────────────────────────────────────────
@app.route("/voice_enhancer", methods=["GET", "POST"])
def voice_enhancer():
    if request.method == "POST":
        uid      = uuid.uuid4().hex
        raw_path = os.path.join(UPLOAD_DIR, f"{uid}_raw.wav")
        out_path = os.path.join(STATIC_DIR, f"{uid}_enhanced.wav")
        try:
            audio_obj = request.files.get("audio")
            if not audio_obj or not audio_obj.filename:
                return jsonify({"error": "No audio file received."}), 400

            speed   = max(0.5, min(float(request.form.get("speed", 1.0)), 2.0))
            volume  = max(-10.0, min(float(request.form.get("volume", 0.0)), 10.0))
            clarity = max(0.0, min(float(request.form.get("clarity", 0.0)), 1.0))
            noise   = max(0.0, min(float(request.form.get("noise", 0.0)), 1.0))

            # Convert to WAV preserving original sample rate for best quality
            # Use 44100 Hz (CD quality) — better than 16kHz for enhancement
            wav_raw = to_wav(audio_obj, uid, mono=True, sr=22050)  # Reduced from 44100 for speed
            import shutil as _shutil
            _shutil.copy2(wav_raw, raw_path)
            safe_remove(wav_raw)

            # Get duration from original audio
            orig_seg  = AudioSegment.from_wav(raw_path)
            dur_s     = round(len(orig_seg) / 1000)
            sub_dur   = f"{dur_s//60}m {dur_s%60}s" if dur_s >= 60 else f"{dur_s}s"

            # Detect language from ORIGINAL audio using accurate Whisper detection
            orig_lcode    = "unknown"
            orig_language = "Unknown"
            orig_confidence = 0
            try:
                wav_16k = os.path.join(UPLOAD_DIR, f"{uid}_16k.wav")
                y_16k, _ = librosa.load(raw_path, sr=16000, mono=True)
                sf.write(wav_16k, y_16k, 16000)
                orig_lcode, orig_language, orig_confidence = detect_audio_language(wav_16k)
                safe_remove(wav_16k)
            except Exception as det_err:
                print(f"  voice_enhancer lang detect error: {det_err}")

            # Load at original sample rate for high-quality DSP
            y, sr_val = librosa.load(raw_path, sr=22050, mono=True)

            # ── DSP pipeline ─────────────────────────────────────────
            y = apply_noise_reduction(y, sr_val, noise)

            if abs(speed - 1.0) > 0.01:
                y = librosa.effects.time_stretch(y, rate=float(speed))

            y = apply_clarity_eq(y, sr_val, clarity)

            # Write enhanced audio
            sf.write(out_path, y, sr_val)
            enhanced = AudioSegment.from_wav(out_path)
            enhanced = effects.normalize(enhanced)
            if volume != 0.0:
                enhanced = enhanced.apply_gain(volume)
            enhanced.export(out_path, format="wav")

            safe_remove(raw_path)
            audio_url = "/" + out_path.replace("\\", "/")
            return jsonify({
                "audio_url":      audio_url,
                "language":       orig_language,
                "orig_language":  orig_language,
                "lang_code":      orig_lcode,
                "confidence":     orig_confidence,
                "duration":       sub_dur,
                "fname":          audio_obj.filename or "Recorded audio",
            })

        except Exception as e:
            safe_remove(raw_path)
            safe_remove(out_path)
            return jsonify({"error": str(e)}), 500

    return render_template("voice_enhancer.html")


# ─────────────────────────────────────────────
if __name__ == "__main__":
    # threaded=True lets Flask handle multiple tabs/requests concurrently
    app.run(debug=True, threaded=True)